import io
import re
from typing import Tuple
from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import UploadFile

from app.config import settings
from app.core.exceptions import ValidationError

MAX_FILE_SIZE = settings.MAX_FILE_SIZE_MB * 1024 * 1024
SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}
MIN_TEXT_LENGTH = 50
MAX_TEXT_LENGTH = 50000


def validate_file(file: UploadFile) -> None:
    if hasattr(file, 'size') and file.size and file.size > MAX_FILE_SIZE:
        raise ValidationError(f"File too large: {file.size / 1024 / 1024:.1f}MB. Max: {settings.MAX_FILE_SIZE_MB}MB")
    if not file.filename:
        raise ValidationError("Filename is required")
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in SUPPORTED_EXTENSIONS:
        raise ValidationError(f"Unsupported file type: {file_ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS.keys())}")


def validate_extracted_text(text: str, filename: str) -> str:
    if not text or not text.strip():
        raise ValidationError(f"No text content found in {filename}")
    cleaned = re.sub(r'\s+', ' ', text.strip())
    if len(cleaned) < MIN_TEXT_LENGTH:
        raise ValidationError(f"Text too short: {len(cleaned)} chars. Min: {MIN_TEXT_LENGTH}")
    if len(cleaned) > MAX_TEXT_LENGTH:
        raise ValidationError(f"Text too long: {len(cleaned)} chars. Max: {MAX_TEXT_LENGTH}")
    return cleaned


def extract_text_from_pdf(file_content: bytes, filename: str) -> str:
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            if not pdf.pages:
                raise ValidationError(f"PDF {filename} contains no pages")
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    for table in page.extract_tables():
                        for row in table:
                            if row:
                                row_text = ' '.join(str(c) for c in row if c)
                                if row_text.strip():
                                    text_parts.append(row_text)
                except Exception as e:
                    print(f"Warning: page {page_num} of {filename}: {e}")
        if not text_parts:
            raise ValidationError(f"No readable text in PDF {filename}")
        return validate_extracted_text('\n'.join(text_parts), filename)
    except ValidationError:
        raise
    except Exception as e:
        raise ValidationError(f"Failed to process PDF {filename}: {str(e)}")


def extract_text_from_docx(file_content: bytes, filename: str) -> str:
    try:
        text_parts = []
        doc = Document(io.BytesIO(file_content))
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))
        if not text_parts:
            raise ValidationError(f"No readable text in DOCX {filename}")
        return validate_extracted_text('\n'.join(text_parts), filename)
    except ValidationError:
        raise
    except Exception as e:
        raise ValidationError(f"Failed to process DOCX {filename}: {str(e)}")


def extract_text_from_txt(file_content: bytes, filename: str) -> str:
    for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
        try:
            return validate_extracted_text(file_content.decode(encoding), filename)
        except UnicodeDecodeError:
            continue
    raise ValidationError(f"Unable to decode {filename}. Please use UTF-8 format.")


async def extract_text_from_file(file: UploadFile) -> Tuple[str, dict]:
    validate_file(file)
    try:
        file_content = await file.read()
        await file.seek(0)
    except Exception as e:
        raise ValidationError(f"Failed to read {file.filename}: {str(e)}")

    file_ext = Path(file.filename).suffix.lower()
    if file_ext == ".pdf":
        extracted_text = extract_text_from_pdf(file_content, file.filename)
    elif file_ext == ".docx":
        extracted_text = extract_text_from_docx(file_content, file.filename)
    elif file_ext == ".txt":
        extracted_text = extract_text_from_txt(file_content, file.filename)
    else:
        raise ValidationError(f"Unsupported file type: {file_ext}")

    return extracted_text, {
        "filename": file.filename,
        "file_type": file_ext,
        "file_size_bytes": len(file_content),
        "content_type": file.content_type,
        "text_length": len(extracted_text),
        "word_count": len(extracted_text.split()),
    }


def get_file_info(file: UploadFile) -> dict:
    file_ext = Path(file.filename).suffix.lower() if file.filename else ""
    return {
        "filename": file.filename,
        "file_type": file_ext,
        "content_type": file.content_type,
        "is_supported": file_ext in SUPPORTED_EXTENSIONS,
        "max_allowed_size_mb": settings.MAX_FILE_SIZE_MB,
    }
